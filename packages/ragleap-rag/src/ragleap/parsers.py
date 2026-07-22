"""
Document text extraction for ragleap-rag.

Core formats (txt, pdf, docx) work with the base install. Everything
else requires the 'formats' extra: pip install ragleap-rag[formats]

Not supported: legacy binary .doc and .ppt (pre-2007 Office formats).
No reliable pure-Python parser exists for these - convert to .docx/
.pptx first (e.g. via LibreOffice headless, or Microsoft Office
"Save As"), or use a dedicated conversion service.
"""
import csv
import io
import json
import logging
import zipfile

logger = logging.getLogger(__name__)

CORE_EXTENSIONS = {".txt", ".pdf", ".docx", ".md"}
EXTENDED_EXTENSIONS = {
    ".xlsx", ".xls", ".csv", ".tsv", ".xsl", ".xslt", ".pptx", ".html", ".htm",
    ".json", ".xml", ".rtf", ".odt", ".ods", ".odp", ".eml", ".zip", ".sql",
    ".epub", ".yaml", ".yml", ".parquet", ".vtt", ".srt",
}
SUPPORTED_EXTENSIONS = CORE_EXTENSIONS | EXTENDED_EXTENSIONS
UNSUPPORTED_LEGACY = {".doc", ".ppt"}


def _require(module_name: str, extra_hint: str = "formats"):
    try:
        return __import__(module_name)
    except ImportError as e:
        raise ValueError(
            f"'{module_name}' is required for this file type — "
            f"pip install ragleap-rag[{extra_hint}]"
        ) from e


def extract_text(filename: str, raw_bytes: bytes) -> str:
    """Extract plain text from raw file bytes, dispatching on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    ext = f".{ext}"

    if ext in UNSUPPORTED_LEGACY:
        raise ValueError(
            f"'{ext}' (legacy binary Office format) is not supported — no reliable "
            f"pure-Python parser exists. Convert to the modern equivalent first "
            f"(.docx/.pptx), e.g. via LibreOffice headless: "
            f"'soffice --headless --convert-to docx yourfile.doc'"
        )

    dispatch = {
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".xls": _extract_xls,
        ".csv": _extract_csv,
        ".tsv": lambda b: _extract_csv(b, delimiter="\t"),
        ".xsl": _extract_xml,
        ".xslt": _extract_xml,
        ".pptx": _extract_pptx,
        ".html": _extract_html,
        ".htm": _extract_html,
        ".json": _extract_json,
        ".xml": _extract_xml,
        ".rtf": _extract_rtf,
        ".odt": _extract_odt,
        ".ods": _extract_ods,
        ".odp": _extract_odp,
        ".eml": _extract_eml,
        ".zip": _extract_zip,
        ".sql": _extract_txt,
        ".epub": _extract_epub,
        ".yaml": _extract_yaml,
        ".yml": _extract_yaml,
        ".parquet": _extract_parquet,
        ".vtt": _extract_subtitle,
        ".srt": _extract_subtitle,
    }

    handler = dispatch.get(ext)
    if handler is None:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}.")

    return handler(raw_bytes)


def _extract_txt(raw_bytes: bytes) -> str:
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed, retrying with latin-1")
        return raw_bytes.decode("latin-1")


def _extract_pdf(raw_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ValueError("pypdf is required for PDF support — pip install ragleap-rag installs it by default") from e

    reader = PdfReader(io.BytesIO(raw_bytes))
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text)
        else:
            logger.warning(f"No extractable text on PDF page {i + 1} (likely scanned/image-only)")
    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        raise ValueError(
            "No text could be extracted from this PDF. It may be a scanned "
            "image-only document, which requires OCR (not currently supported)."
        )
    return full_text


def _extract_docx(raw_bytes: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise ValueError("python-docx is required for DOCX support — pip install ragleap-rag installs it by default") from e

    document = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        raise ValueError("No text could be extracted from this DOCX file — it may be empty.")
    return full_text


def _extract_xlsx(raw_bytes: bytes) -> str:
    openpyxl = _require("openpyxl")
    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(c) for c in row if c is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_xls(raw_bytes: bytes) -> str:
    xlrd = _require("xlrd")
    wb = xlrd.open_workbook(file_contents=raw_bytes)
    parts = []
    for sheet in wb.sheets():
        parts.append(f"[Sheet: {sheet.name}]")
        for row_idx in range(sheet.nrows):
            row = sheet.row_values(row_idx)
            row_text = "\t".join(str(c) for c in row if c != "")
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_csv(raw_bytes: bytes, delimiter: str = ",") -> str:
    text = _extract_txt(raw_bytes)
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    return "\n".join("\t".join(row) for row in reader)


def _extract_pptx(raw_bytes: bytes) -> str:
    pptx = _require("pptx", extra_hint="formats")
    prs = pptx.Presentation(io.BytesIO(raw_bytes))
    parts = []
    for i, slide in enumerate(prs.slides, start=1):
        slide_text = []
        for shape in slide.shapes:
            if shape.has_text_frame:
                for para in shape.text_frame.paragraphs:
                    text = "".join(run.text for run in para.runs)
                    if text.strip():
                        slide_text.append(text)
        if slide.has_notes_slide and slide.notes_slide.notes_text_frame:
            notes = slide.notes_slide.notes_text_frame.text
            if notes.strip():
                slide_text.append(f"[Speaker notes: {notes}]")
        if slide_text:
            parts.append(f"[Slide {i}]\n" + "\n".join(slide_text))
    return "\n\n".join(parts)


def _extract_html(raw_bytes: bytes) -> str:
    bs4 = _require("bs4")
    soup = bs4.BeautifulSoup(raw_bytes, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _extract_json(raw_bytes: bytes) -> str:
    data = json.loads(_extract_txt(raw_bytes))
    return json.dumps(data, indent=2, ensure_ascii=False)


def _extract_xml(raw_bytes: bytes) -> str:
    bs4 = _require("bs4")
    soup = bs4.BeautifulSoup(raw_bytes, "xml")
    return soup.get_text(separator="\n", strip=True)


def _extract_rtf(raw_bytes: bytes) -> str:
    striprtf = _require("striprtf")
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(_extract_txt(raw_bytes))


def _extract_odt(raw_bytes: bytes) -> str:
    odf_opendoc = _require("odf.opendocument", extra_hint="formats")
    from odf import text as odf_text
    from odf.opendocument import load
    doc = load(io.BytesIO(raw_bytes))
    paragraphs = doc.getElementsByType(odf_text.P)
    return "\n".join(str(p) for p in paragraphs)


def _extract_ods(raw_bytes: bytes) -> str:
    from odf.opendocument import load
    from odf.table import Table, TableRow, TableCell
    doc = load(io.BytesIO(raw_bytes))
    parts = []
    for table in doc.getElementsByType(Table):
        for row in table.getElementsByType(TableRow):
            cells = row.getElementsByType(TableCell)
            row_text = "\t".join(str(c) for c in cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_odp(raw_bytes: bytes) -> str:
    from odf import text as odf_text
    from odf.opendocument import load
    doc = load(io.BytesIO(raw_bytes))
    paragraphs = doc.getElementsByType(odf_text.P)
    return "\n".join(str(p) for p in paragraphs)


def _extract_eml(raw_bytes: bytes) -> str:
    import email
    from email import policy
    msg = email.message_from_bytes(raw_bytes, policy=policy.default)
    parts = [f"From: {msg.get('From', '')}", f"To: {msg.get('To', '')}", f"Subject: {msg.get('Subject', '')}"]
    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        content = body.get_content()
        if body.get_content_type() == "text/html":
            content = _extract_html(content.encode("utf-8"))
        parts.append(content)
    return "\n".join(parts)


def _extract_zip(raw_bytes: bytes) -> str:
    """Extracts and concatenates text from every supported file inside the zip."""
    parts = []
    with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
        for name in z.namelist():
            if name.endswith("/"):
                continue
            ext = f".{name.lower().rsplit('.', 1)[-1]}" if "." in name else ""
            if ext not in SUPPORTED_EXTENSIONS or ext == ".zip":
                continue
            try:
                inner_bytes = z.read(name)
                inner_text = extract_text(name, inner_bytes)
                if inner_text.strip():
                    parts.append(f"[File: {name}]\n{inner_text}")
            except Exception as e:
                logger.warning(f"Skipping '{name}' inside zip — extraction failed: {e}")
    if not parts:
        raise ValueError("No extractable text found in any file inside this zip.")
    return "\n\n".join(parts)


def _extract_epub(raw_bytes: bytes) -> str:
    ebooklib = _require("ebooklib")
    import ebooklib as _eb
    from ebooklib import epub
    import tempfile, os as _os
    with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
        f.write(raw_bytes)
        tmp_path = f.name
    try:
        book = epub.read_epub(tmp_path)
        parts = []
        for item in book.get_items_of_type(_eb.ITEM_DOCUMENT):
            html_text = _extract_html(item.get_content())
            if html_text.strip():
                parts.append(html_text)
        return "\n\n".join(parts)
    finally:
        _os.unlink(tmp_path)


def _extract_yaml(raw_bytes: bytes) -> str:
    yaml = _require("yaml", extra_hint="formats")
    data = yaml.safe_load(_extract_txt(raw_bytes))
    return json.dumps(data, indent=2, ensure_ascii=False, default=str)


def _extract_parquet(raw_bytes: bytes) -> str:
    pyarrow = _require("pyarrow")
    import pyarrow.parquet as pq
    table = pq.read_table(io.BytesIO(raw_bytes))
    df = table.to_pandas()
    return df.to_csv(index=False, sep="\t")


def _extract_subtitle(raw_bytes: bytes) -> str:
    """Strips timestamps/cue numbers from .vtt/.srt subtitle files, keeping spoken text."""
    text = _extract_txt(raw_bytes)
    lines = text.splitlines()
    kept = []
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == "WEBVTT":
            continue
        if stripped.isdigit():
            continue
        if "-->" in stripped:
            continue
        kept.append(stripped)
    return "\n".join(kept)
